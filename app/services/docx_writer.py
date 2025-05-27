from docx import Document

def write_meta_raw_docx(about_text, financials_text, output_path):
    doc = Document()
    doc.add_heading("Company Overview", 1)
    doc.add_paragraph(about_text.strip())
    doc.add_page_break()
    doc.add_heading("Company Financials", 1)
    doc.add_paragraph(financials_text.strip())
    doc.save(output_path)

def write_meta_links_docx(about_links, financial_links, news_links, output_path):
    doc = Document()
    doc.add_heading("Links – Company Overview", 1)
    for link in about_links:
        doc.add_paragraph(link, style="List Bullet")

    doc.add_page_break()
    doc.add_heading("Links – Company Financials", 1)
    for link in financial_links:
        doc.add_paragraph(link, style="List Bullet")

    if news_links:
        doc.add_page_break()
        doc.add_heading("Links – News Sources", 1)
        for link in news_links:
            doc.add_paragraph(link, style="List Bullet")

    doc.save(output_path)

def write_compiled_docx(summary_text, output_path):
    doc = Document()
    doc.add_heading("Company Profile Summary", 1)
    doc.add_paragraph(summary_text.strip())
    doc.save(output_path)
